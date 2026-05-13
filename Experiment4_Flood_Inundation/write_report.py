"""Generate Word report summarizing the flood analysis workflow."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

OUT_DIR = Path(__file__).resolve().parent
IMG_DIR = OUT_DIR

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
style.font.size = Pt(11)
style.font.name = "Calibri"

# --- Title ---
title = doc.add_heading("Flood Inundation Analysis — Workflow Report", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# --- 1. DEM Generation ---
doc.add_heading("1. Synthetic DEM Generation", level=1)
doc.add_paragraph(
    "A 100\xd7100 synthetic Digital Elevation Model was generated using Fractional Brownian Motion (FBM) "
    "with 6 octaves to produce realistic terrain features. A gentle sinusoidal ridge was added to mimic "
    "natural topography."
)
doc.add_paragraph(
    f"Elevation range: 30.00 – 80.00 m  |  Mean: 53.99 m  |  Grid: 100\xd7100 cells"
)
doc.add_paragraph("Method: fbm_noise() with octaves=6, lacunarity=2.0, gain=0.5, seed=42")
doc.add_paragraph("Files: generate_dem.py, dem_synthetic_100x100.npy, dem_synthetic_100x100.csv")

# --- 2. Inundation Computation ---
doc.add_heading("2. Flood Inundation Computation", level=1)
doc.add_paragraph("Three core functions compute inundation from a DEM given a water level (WL):")
doc.add_paragraph("flood_mask(dem, WL): Boolean mask — True where elevation < WL", style="List Bullet")
doc.add_paragraph("flood_depth(dem, WL): Depth = WL – elevation for flooded cells, 0 elsewhere", style="List Bullet")
doc.add_paragraph("flood_area_pct(mask): Flooded cells / total cells x 100%", style="List Bullet")
doc.add_paragraph("Convenience wrapper flood_stats() returns all metrics in one call.")
doc.add_paragraph("")

# Table: sample results
table = doc.add_table(rows=8, cols=5, style="Light Grid Accent 1")
headers = ["Water Level (m)", "Area (%)", "Mean Depth (m)", "Max Depth (m)", "Volume (m\xb3)"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data_rows = [
    ("40", "3.79", "2.23", "10.00", "845"),
    ("45", "14.48", "3.47", "15.00", "5,019"),
    ("50", "32.49", "5.07", "20.00", "16,467"),
    ("55", "55.04", "6.95", "25.00", "38,233"),
    ("60", "75.83", "9.37", "30.00", "71,043"),
    ("65", "90.60", "12.47", "35.00", "112,967"),
    ("70", "98.15", "16.35", "40.00", "160,523"),
]
for r, row_data in enumerate(data_rows):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val
doc.add_paragraph("Table: Flood statistics across water levels").italic = True
doc.add_paragraph("")

# --- 3. Visualization ---
doc.add_heading("3. Visualization", level=1)
doc.add_paragraph(
    "A 3x3 grid figure was generated comparing water levels 45, 55, and 65 m across three rows: "
    "(a) raw DEM in terrain colormap, (b) flood extent overlay (blue mask on grayscale DEM), "
    "and (c) flood depth heatmap (Blues colormap). Each subplot includes a color bar and title."
)
img_path = IMG_DIR / "flood_inundation_plot.png"
if img_path.exists():
    doc.add_picture(str(img_path), width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Figure 1: Flood inundation visualization at WL = 45, 55, 65 m").italic = True
doc.add_paragraph("")

# --- 4. Trend Analysis ---
doc.add_heading("4. Rising-Water Trend Analysis", level=1)
doc.add_paragraph(
    "Water level was incremented from 40.0 m to 50.0 m in 0.5 m steps. At each step, the "
    "flooded area percentage was recorded and plotted."
)
img_path2 = IMG_DIR / "flood_trend_curve.png"
if img_path2.exists():
    doc.add_picture(str(img_path2), width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Figure 2: Water level vs. Inundation area curve").italic = True
doc.add_paragraph(
    "Key finding: area increased from 3.79% to 32.49% (8.6x) over a 10 m rise. "
    "A pronounced inflection between 47–48 m indicates a flat terrain band that is "
    "highly sensitive to small water-level changes."
)
doc.add_paragraph("")

# --- 5. Validation ---
doc.add_heading("5. Physical-Sense Validation", level=1)
doc.add_paragraph("Five validation checks were performed across 101 water levels (30–80 m, step 0.5 m):")
checks_list = [
    "Area percentage always in [0, 100]",
    "Maximum depth = water level – minimum elevation (when flooded)",
    "All depth values are non-negative",
    "Depth > 0 exactly where mask is True (mask/depth consistency)",
    "Depth values equal water level – DEM for all flooded cells",
]
for ch in checks_list:
    doc.add_paragraph(ch, style="List Bullet")
doc.add_paragraph(
    "\nMonotonicity check: flood area must be non-decreasing with increasing water level."
)
doc.add_paragraph("")
doc.add_paragraph("Result: ALL 506 checks PASSED (101 levels x 5 + 1 monotonicity).")
doc.add_paragraph("")

# --- 6. Files ---
doc.add_heading("6. Output Files", level=1)
files_list = [
    ("generate_dem.py", "Synthetic DEM generator (FBM noise)"),
    ("dem_synthetic_100x100.npy", "DEM data (NumPy binary, float32)"),
    ("dem_synthetic_100x100.csv", "DEM data (CSV, 100x100 grid)"),
    ("flood_inundation.py", "Flood computation functions"),
    ("visualize_flood.py", "Matplotlib visualization (3x3 grid)"),
    ("flood_inundation_plot.png", "Visualization output (419 KB, 150 DPI)"),
    ("flood_trend.py", "Water-level vs. area curve analysis"),
    ("flood_trend_curve.png", "Trend curve output"),
    ("validate_flood.py", "Physical-sense validation suite"),
    ("validation_results.npz", "Validation data archive"),
    ("write_report.py", "This report generator"),
]
for fname, desc in files_list:
    doc.add_paragraph(f"{fname} — {desc}", style="List Bullet")

doc.add_paragraph("")
doc.add_paragraph("All validation passed. Results are physically consistent.", style="Intense Quote")

# --- Save ---
report_path = OUT_DIR / "Flood_Inundation_Analysis_Report.docx"
doc.save(report_path)
print(f"Report saved: {report_path}")
